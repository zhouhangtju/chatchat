import docx
from collections import defaultdict
from docx.table import Table
from docx.oxml.shared import qn
import base64
from langchain.docstore.document import Document
from configs import CHUNK_SIZE

class MyDocxLoader:
    def __init__(self, file_path, chunk_size=CHUNK_SIZE):
        self.file_path = file_path
        self.chunk_size = chunk_size

    def __get_document_title(self, doc):
        title = doc.core_properties.title
        if not title:
            if doc.paragraphs:
                first_paragraph = doc.paragraphs[0]
                if first_paragraph.style.name.lower() in ['title', 'heading 1']:
                    title = first_paragraph.text
        if not title:
            title = ""
        return title

    def __extract_docx_content(self):
        doc = docx.Document(self.file_path)
        content = defaultdict(list)
        document_title = self.__get_document_title(doc)
        heading_stack = [document_title]
        current_heading = document_title

        for element in doc.element.body:
            if element.tag.endswith('p'):
                paragraph = docx.text.paragraph.Paragraph(element, doc)
                if paragraph.style and paragraph.style.name.startswith('Heading'):
                    heading_level = int(paragraph.style.name[-1])

                    while len(heading_stack) > heading_level:
                        heading_stack.pop()

                    if len(heading_stack) < heading_level:
                        heading_stack.extend([""] * (heading_level - len(heading_stack)))

                    heading_stack[heading_level - 1] = paragraph.text
                    current_heading = document_title + " > "+" > ".join(filter(None, heading_stack))
                else:
                    content[current_heading].append({"type": "text", "content": paragraph.text})
            elif element.tag.endswith('tbl'):
                table = Table(element, doc)
                table_data = []
                for row in table.rows:
                    table_data.append([cell.text for cell in row.cells])
                content[current_heading].append({"type": "table", "content": table_data})
            elif element.tag.endswith('r'):
                for run in element.findall(".//"+qn('w:drawing')):
                    blip = run.find(".//"+qn('a:blip'))
                    if blip is not None:
                        image_rid = blip.get(qn('r:embed'))
                        image_part = doc.part.related_parts[image_rid]
                        image_bytes = image_part.blob
                        image_base64 = base64.b64encode(image_bytes).decode()
                        content[current_heading].append({"type": "image", "content": image_base64})
        return content

    def __format_output(self, content):
        output = []
        for heading, elements in content.items():
            block = {
                "metadata": {"title": heading},
                "content": []
            }
            for element in elements:
                if element['type'] == 'text':
                    block['content'].append(f"{element['content']}")
                elif element['type'] == 'table':
                    table_str = "\n".join([", ".join(row) for row in element['content']])
                    block['content'].append(table_str)
                elif element['type'] == 'image':
                    block['content'].append(f"[Base64 encoded image data]")
            block['content'] = "\n".join(block['content'])

            if len(block['content']) > self.chunk_size:
                chunks = [block['content'][i:i+self.chunk_size] for i in range(0, len(block['content']), self.chunk_size)]
                for chunk in chunks:
                    output.append({"metadata": {"title": heading}, "content": chunk})
            else:
                output.append(block)
        return output

    def load(self):
        extracted_content = self.__extract_docx_content()
        formatted_output = self.__format_output(extracted_content)
        docs = [Document(page_content=block['content'], metadata=block['metadata']) for block in formatted_output]
        return docs

# Example usage
if __name__ == "__main__":
    loader = MyDocxLoader("./docx-test/中国移动集中网络云承载网网管&业支CE局数据配置规范初稿（华为）V4.0.docx")
    documents = loader.load()
    for doc in documents:
        print(doc)
        # print(f"Title: {doc.metadata['title']}")
        # print(f"Content: {doc.page_content}")
        print("#" * 100)
