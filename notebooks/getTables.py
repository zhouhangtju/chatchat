import pandas as pd
import os
import docx
import shutil
from PIL import Image
from io import BytesIO

from docx.document import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.oxml.section import CT_SectPr
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from docx.section import Section

# from docsplitter.getPics import tbl2_text_withpics
from utils import *


# IMAGE_LABEL_TEMP = "#<!image:{}>#"
# from configs import *


IMAGE_LABEL_TEMP = "#<!image:{}>#"
def iter_block_items(parent):
    """
    Yield each paragraph and table child within *parent*, in document order.
    Each returned value is an instance of either Table or Paragraph. *parent*
    would most commonly be a reference to a main Document object, but
    also works for a _Cell object, which itself can contain paragraphs and tables.
    """
    if isinstance(parent, Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")
    blocks = []
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            # yield Paragraph(child, parent)
            blocks.append(Paragraph(child, parent))
        elif isinstance(child, CT_Tbl):
            # yield Table(child, parent)
            blocks.append(Table(child, parent))
        # elif isinstance(child, CT_SectPr):
        # #     blocks.append(CT_SectPr)
        #     blocks.append(child)
        # adding picture
        # elif isinstance(child, CT_Picture):
        #     blocks.append(Image(child, parent))

    return blocks


def save_tbl2df(table: Table, save_to=".file1.csv"):
    data = []
    for row in table.rows:
        row_data = []
        for cell in row.cells:
            if cell.text:
                row_data.append("\t" + cell.text)
        data.append(row_data)

    df = pd.DataFrame(data)
    df.columns = df.iloc[0]  # Set column names to the first row's values
    df = df.iloc[1:]  # Remove the first row
    df.to_csv(save_to)
    # df.to_excel(save_to)
    return


def tbl2_markdown(table: Table) -> str:
    data = []
    for row in table.rows:
        row_data = []
        for cell in row.cells:
            if cell.text:
                row_data.append(cell.text)
        data.append(row_data)

    df = pd.DataFrame(data)
    df.columns = df.iloc[0]  # Set column names to the first row's values
    df = df.iloc[1:]  # Remove the first row
    return df.to_markdown()

# def save_pics(pic, pic_id, pics_folder, get_image_ocr=False) -> str:
#     """保存图片 并返回图片标记，或者图片内容"""
#     try:
#         ext = pic.ext  # 后缀
#         blob = pic.blob  # 二进制内容
#         pil_img = Image.open(BytesIO(blob))
#         # img_label = f'#<!image:image img_{pic_id}.{ext}>#'

#         img_name = f"img_{pic_id}.{ext}"
#         pic_save_path = os.path.join(pics_folder, img_name)
#         pil_img.save(fp=pic_save_path)
#         # new_doc.paragraphs[-1].add_run(img_label)
#         if get_image_ocr:
#             format = ext
#             pic_data = open(pic_save_path, "rb").read()
#             img_text = get_images_ocr(pic_data=pic_data, format=format)
#             # page_content.append(IMAGE_LABEL_TEMP.format(img_name))
#             # page_content.append("".join(img_text))
#             img_text = "".join(img_text)
#             return img_name+" "+img_text
#         return img_name
#     except AttributeError:
#         return ""


def get_docx_imgs_totext(pic, pic_id, pics_folder, get_image_ocr=False,
                         **kwargs) -> str:
    """保存图片 并返回图片标记，或者图片内容"""
    try:
        ext = pic.ext  # 后缀
        blob = pic.blob  # 二进制内容
        pil_img = Image.open(BytesIO(blob))
        # img_label = f'#<!image:image img_{pic_id}.{ext}>#'

        img_name = f"img_{pic_id}.{ext}"
        pic_save_path = os.path.join(pics_folder, img_name)
        pil_img.save(fp=pic_save_path)
        if get_image_ocr:
            format = ext
            pic_data = open(pic_save_path, "rb").read()
            img_text = get_images_ocr(pic_data=pic_data, format=format)
            # page_content.append(IMAGE_LABEL_TEMP.format(img_name))
            # page_content.append("".join(img_text))
            img_text = "".join(img_text)
            return IMAGE_LABEL_TEMP.format(img_name)+img_text
        # new_doc.paragraphs[-1].add_run(img_label)
        return IMAGE_LABEL_TEMP.format(img_name)
    
    except AttributeError:
        return ""








# def delete_paragraph(paragraph):
#     p = paragraph._element
#     p.getparent().remove(p)
#     # p._p = p._element = None
#     paragraph._p = paragraph._element = None


# if __name__=="__main__":
#     path = "../test docs/奥体场馆群-杭州奥体中心主体育场-一馆一案.docx"
#     doc = docx.Document(path)
#     blocks = iter_block_items(doc)
#     if os.path.exists("./pics_tmp"):
#         shutil.rmtree("./pics_tmp")
#         os.mkdir("./pics_tmp")
#     else:
#         os.mkdir("./pics_tmp")

#     i=0
#     for block in blocks:
#         if isinstance(block, Table):
#             tmp = tbl2_text_withpics(doc, block, pics_folder=f"./pics_tmp")
#             i += 1
#             print(tmp)

# if __name__ == "__main__":
#     path = "test docs\API编排集群部署手册.docx"
#     doc = docx.Document(path)
#     blocks = iter_block_items(doc)
#     for block in blocks:
#         if isinstance(block, Section):
#             print(block)

#     path = "亚运知识库/保障一馆一册/保障一馆一册/杭州2022年第19届亚运会宁波象山沙滩排球场馆通信网络服务保障一馆一册V1.0-20230321定稿.docx"
#     # path = "沙滩排球馆数据.docx"
#     doc = docx.Document(path)
#     paras_ = iter_block_items(doc)
#
#     new_doc = doc
#     print(len(new_doc.tables))
#
#     last_block = None
#     if os.path.exists("./tables"):
#         shutil.rmtree("./tables")
#         os.mkdir("./tables")
#     else:
#         os.mkdir("./tables")
#
#     if os.path.exists("./pics"):
#         shutil.rmtree("./pics")
#         os.mkdir("./pics")
#     else:
#         os.mkdir("./pics")
#
#     tbl_idx = 0
#     for i, block in enumerate(paras_):
#
#         if isinstance(block, Paragraph):
#             pass
#         elif isinstance(block, Table):
#
#             if isinstance(last_block, Paragraph) and \
#                     last_block.text.startswith("表"):
#                 tbl_name = last_block.text
#             elif isinstance(paras_[i + 1], Paragraph) and \
#                     paras_[i + 1].text.startswith("表"):
#                 tbl_name = paras_[i + 1].text
#             else:
#                 tbl_name = tbl_idx
#
#             save_tbl2df(block, save_to=f'./tables/table_{tbl_name}.csv')
#             delete_paragraph(new_doc.tables[0])
#             tbl_markdown = tbl2_markdown(block)
#             new_doc.paragraphs[i].insert_paragraph_before(tbl_markdown)
#             tbl_idx += 1
#         last_block = block
#
#     # print(len(new_doc.paragraphs))
#     # print(len(paras_))
#
#     new_path = "./new_doc.docx"
#     # new_path = path.replace(".docx", "_tbl2mkdn.docx")
#     new_doc.save(new_path)
