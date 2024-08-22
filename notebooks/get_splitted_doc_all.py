'''
Author       : your name
Date         : 2024-07-31 07:08:10
LastEditors  : your name
LastEditTime : 2024-08-01 06:45:04
FilePath     : /get_splitted_doc_all.py
Description  : 
Copyright 2024 OBKoro1, All Rights Reserved. 
2024-07-31 07:08:10
'''
import sys
from pathlib import Path
sys.path.append(str(Path(__file__).parent))
sys.path.append('./getTables.py')

import os
import docx
import traceback
import shutil
import re

# import pandas as pd
from PIL import Image
from io import BytesIO

# import matplotlib.pyplot as plt

from docx.text.paragraph import Paragraph
from docx.table import Table

from getTables import iter_block_items,get_docx_imgs_totext
from utils import split_text_by_sentense,clean_spaces,get_toc_hierarchy,get_tokens,delete_paragraph
from docx_extract import get_toc_docx
from getPics import get_picture_list,tbl2_text_withpics



# if __name__ == "__main__":
#     from docx_extract import *
#     from pdf_extract2 import *
# else:
#     from docsplitter.docx_extract import *
#     from docsplitter.pdf_extract2 import *


# IMAGE_LABEL_TEMP = "#<!image:{}>#"
# from configs import *

# pdf预处理参数
headers_height = 50
footers_height = 770


def completely_rm_space(t):
    """完全删除 空格和制表符"""
    text_no_sp = t.replace(" ", "")
    text_no_sp = re.sub(r"\s+", "", text_no_sp)
    text_no_sp = re.sub(r"\t+", "", text_no_sp).strip()
    return text_no_sp


def docx_general_preprocess(
    path,
    pics_folder: str,
    **kwargs,
):
    """
    做如下预处理：
    保存图片 在原文中定位
    转化表格，保存为excel并转文字 在原文中定位？ # 这个先不做
    输出：list[list]: docx每一段[doc每一句+图片 or docx表格]
    """
    # print("process docx", get_tables)
    get_tables = kwargs.get("get_tables", False)
    # get_image_ocr = kwargs.get("get_image_ocr", False)
    if not path.endswith(".docx"):
        raise ValueError("文件类型必须是docx")
        # return False
    filename = os.path.split(path)[-1]  # 有后缀文件名

    total_docx_content = []
    total_image_para = {}  # 图片名称+段落号

    # 打开docx文件
    doc = docx.Document(path)
    blocks = iter_block_items(doc)

    pic_count = 0
    try:
        for i, block in enumerate(blocks):
            paragraph_content = []
            para_imgs = []
            p = block

            if isinstance(block, Paragraph):
                # print(p.text)
                # p_text = p.text.replace("\n", "  ")
                p_text = p.text
                # if p.text:
                texts = split_text_by_sentense(p_text)
                for t in texts:
                    if t.strip():
                        t = clean_spaces(t)
                        paragraph_content.append(t)
                # print(paragraph_content)
                # 获取本段图片
                images = get_picture_list(doc, [p])
                if images:
                    for image in images:
                        image_saved_name = get_docx_imgs_totext(
                            image, pic_count, pics_folder, **kwargs
                        )
                        # print(image_saved_name)
                        if image_saved_name:
                            # new_doc.paragraphs[-1].add_run(image_saved_label)
                            # para_imgs.append(IMAGE_LABEL_TEMP.format(image_saved_name))
                            paragraph_content.append(image_saved_name)
                            pic_count += 1
                        # print(para_imgs)
                # if not p.text.strip() and not para_imgs:
                #     continue
                # insert_text_by_toc_with_update(p.text)
            elif isinstance(block, Table):
                if get_tables:
                    # print(1)
                    # 提取表格文字和图片，转换成text
                    tbl_text, tbl_pic_count = tbl2_text_withpics(
                        doc, block, pics_folder=pics_folder
                    )
                    pic_count += tbl_pic_count
                    # print(tbl_text)
                    # new_doc.paragraphs[-1].add_run("; ".join(tbl_text))
                    paragraph_content = tbl_text
            # print(para_imgs)
            total_docx_content.append(paragraph_content)
            if para_imgs:
                total_image_para[i] = para_imgs
        # print(total_image_para)
        return total_docx_content, total_image_para
    except Exception as e:
        traceback.print_exc()
        err_msg = str(e)
        print(f"文档预处理失败！Error message: {err_msg}")
        return [], {}



def insert_text_by_toc(text, curr_images):
    """
    根据toc插入文字
    """
    global new_doc, title_idx, toc, toc_long
    if title_idx < len(toc):
        curr_toc = toc[title_idx][-1].strip()
    # 去除伪标题
    if title_idx < len(toc) and curr_toc == "DummyTitle":
        title_idx += 1

    p_text_no_sp = completely_rm_space(text)
    # 增加新标题，新的一段
    if title_idx < len(toc) and (
        p_text_no_sp.lower() == completely_rm_space("".join(toc[title_idx]).lower())
        or p_text_no_sp.lower() == completely_rm_space(curr_toc).lower()
    ):
        # print(toc[title_idx], text)
        # 检查是否上一段落为空，如果是则删除该段落
        if len(new_doc.paragraphs) >= 2:
            if (
                new_doc.paragraphs[-1].runs[-1].font.bold
                or not new_doc.paragraphs[-1].text
            ):
                # print(new_doc.paragraphs[-1].text)
                delete_paragraph(new_doc.paragraphs[-1])

        # # 在上一段罗结尾加上图片
        # new_doc.paragraphs[-1].add_run("; ".join(curr_images))
        # 在上一段落结尾加上句号
        if not new_doc.paragraphs[-1].text.strip().endswith("。"):
            # print(new_doc.paragraphs[-1].runs[-1].text)
            # if new_doc.paragraphs[-1].runs[-1].text.strip():
            new_doc.paragraphs[-1].add_run("。")

        paragraph1 = new_doc.add_paragraph()
        toc_title = toc_long[title_idx]
        if not toc_title.endswith(
            (
                "。",
                "；",
                "，",
                "：",
                ";",
            )
        ):
            toc_title = toc_title.strip() + "："
        run1 = paragraph1.add_run(toc_title)

        run1.font.bold = True
        title_idx += 1
    # 不另起一段，直接在本段内加内容
    else:
        
        p_text = text.strip().strip("\t")


        if not p_text:
            return
        if p_text:
            new_doc.paragraphs[-1].add_run(p_text)


def doc_split_main(**params):
    global new_doc
    # save_ext = params.get("split_mode")
    new_doc = docx.Document()

    # 添加文本描述
    if params.get("add_desc", ""):
        new_doc.add_paragraph("")
        run1 = new_doc.paragraphs[-1].add_run("文档总结：")
        run1.font.bold = True
        new_doc.paragraphs[-1].add_run(params.get("add_desc"))
    new_doc.add_paragraph("")

    pics_folder = params.get("pics_folder")
    # 创建文件夹
    if os.path.exists(pics_folder):
        shutil.rmtree(pics_folder)
        os.mkdir(pics_folder)
    else:
        os.mkdir(pics_folder)
    print(f"pictures will be saved to: {pics_folder}")

    path = params.get("path")
    file_type = path.split(".")[-1]
    # 文档预处理
    try:
        if file_type == "pdf":
            pass
            # doc_processed, image_dict = pdf_general_preprocess(**params)
        else:   # file_type == "docx":
            doc_processed, image_dict = docx_general_preprocess(**params)
        # else:
        #     raise ValueError("不支持此类文档！请上传pdf或者docx进行切分。")
    except Exception as e:
        raise ValueError("文件预处理失败，"+str(e))

    # # 决定切分模式
    by_title = params.get("by_title")
    if by_title:
        global toc
        if file_type == "pdf":
            pass
            # toc = get_toc_pdf(path)
        else:   # file_type == "docx":
            toc = get_toc_docx(path)
        # 如果获取到目录则继续
        if toc:
            doc_splitted = split_by_toc(doc_processed, image_dict, **params)
        # 没有获取到目录, 报错
        else:
            raise ValueError("无法从文档中提取到目录！请调整切分方法。")
    else:
        doc_splitted = split_by_tokens(doc_processed, image_dict, **params)
        
    # 切分失败
    if not doc_splitted:
        raise ValueError("文档切分失败！")

    return doc_splitted


def split_by_toc(
    doc_processed, image_dict, path: str, pics_folder: str, new_path: str, **kwargs
):
    # print(get_tables)
    global new_doc
    global toc, toc_long, title_idx, curr_images
    file_type = path.split(".")[-1]

    filename = os.path.split(path)[-1]  # 有后缀文件名

    # 获取合并标题：文件名+一级+二级+三级标题
    toc_long = get_toc_hierarchy(
        toc, with_idx=False, file_title=filename.strip("." + file_type)
    )
    try:
        title_idx = 0
        for i, parg in enumerate(doc_processed):
            curr_images = image_dict.get(i, [])
            for line in parg:
                if line.strip() or curr_images:
                    insert_text_by_toc(text=line.strip(), curr_images=curr_images)
        # new_doc.paragraphs[-1].add_run("。")
        new_doc.save(new_path)
        return new_path
    except Exception as e:
        traceback.print_exc()
        print(f"文档分段失败！Error message: {str(e)}")
        return ""


def split_by_tokens2(
    doc_processed,
    image_dict,
    path: str,
    new_path: str = "",
    split_size: int = 500,
    **kwargs,
):
    global new_doc

    file_type = path.split(".")[-1]
    curr_images = []
    try:
        for i, parg in enumerate(doc_processed):
            parg_content = "".join(parg).strip().strip("\t")
            if not parg_content:
                continue
            elif get_tokens(parg_content) <= split_size:
                new_doc.add_paragraph("".join(parg))
                new_doc.add_paragraph("")
            else:
                for line in parg:
                    if (
                        get_tokens(new_doc.paragraphs[-1].text) + get_tokens(line) + 1
                        > split_size
                    ):
                        # if not new_doc.paragraphs[-1].runs[-1].text.endswith("。"):
                        #     new_doc.paragraphs[-1].add_run("。")
                        if curr_images:
                            curr_images += image_dict.get(i, [])
                            new_doc.paragraphs[-1].add_run("; ".join(set(curr_images)))

                        if line.strip("\n").strip():
                            new_doc.add_paragraph(line)
                    else:
                        # if new_doc.paragraphs[-1].text:
                        #     new_doc.paragraphs[-1].add_run("  ")
                        if line.strip("\n").strip():
                            new_doc.paragraphs[-1].add_run(line)
                        curr_images = image_dict.get(i, [])
        # new_doc.paragraphs[-1].add_run("。")
        new_doc.save(new_path)
        return new_path
    except Exception as e:
        traceback.print_exc()
        print(f"文档分段失败！Error message: {str(e)}")
        return ""


def split_by_tokens(
    doc_processed,
    image_dict,
    path: str,
    pics_folder: str = "",
    new_path: str = "",
    split_size: int = 500,
    file_type: str = "pdf",
    **kwargs,
):
    global new_doc
    print("文档预处理完成！", flush=True)
    curr_images = []
    try:
        for i, page in enumerate(doc_processed):
            for line in page:
                if len(new_doc.paragraphs[-1].text) + len(line) + 1 > split_size:
                    if not new_doc.paragraphs[-1].text.strip().endswith("。"):
                        # if new_doc.paragraphs:
                        new_doc.paragraphs[-1].add_run("。")
                    if curr_images:
                        curr_images += image_dict.get(i, [])
                        new_doc.paragraphs[-1].add_run("；".join(set(curr_images)))

                    if line.strip("\n").strip():
                        new_doc.add_paragraph(line)
                else:
                    if new_doc.paragraphs[-1].text:
                        new_doc.paragraphs[-1].add_run("  ")
                    if line.strip("\n").strip():
                        new_doc.paragraphs[-1].add_run(line)
                    curr_images = image_dict.get(i, [])
            # curr_images += image_dict.get(i, [])
        # new_doc.paragraphs[-1].add_run("。")
        new_doc.save(new_path)
        return new_path
    except Exception as e:
        traceback.print_exc()
        print(f"文档分段失败！Error message: {str(e)}")
        return ""


if __name__ == "__main__":
    doc_split_main(pics_folder="./picture", \
                   path="/root/llm/Langchain-Chatchat-master/notebooks/docx-test/中国移动集中网络云承载网网管&业支CE局数据配置规范初稿（华为）V4.0.docx", \
                   by_title=True,
                   new_path="./new_path.txt")